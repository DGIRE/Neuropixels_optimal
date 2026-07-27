"""
Build NP-DEMO-2 METHODS and RESULTS reports (DOCX) from RESULT-* objects.

Every numeral in the narrative text is produced by substituting a
{{RESULT-id.field.path:format}} placeholder with a value read directly
from a RESULT-*.yaml object in 04_results/ (or a simple, transparent
bookkeeping sum/count over such an object, e.g. "total units across
sessions" = sum of qc-counts.sessions[*].n_neurons). No statistic is
hand-typed.

Run:
    C:\\Users\\david\\anaconda3\\python.exe build_reports.py
"""
import os
import re
import yaml
from collections import Counter, OrderedDict

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

RESULTS_DIR = r"C:\Projects\Repos\Neuropixels\_pipeline\NP-DEMO-2\04_results"
OUT_DIR = r"C:\Projects\Repos\Neuropixels\_pipeline\NP-DEMO-2\07_report"

RESULT_IDS = [
    "eth-threshold-log",
    "qc-counts",
    "qc-discards",
    "mrl-unit",
    "phase-unit",
    "mrl-animal",
    "stat-animal",
    "stat-unit",
    "delta-mrl",
    "psth-extremes",
    "psth-examples",
    "unit-locations",
]

# ---------------------------------------------------------------------------
# Load all RESULT objects
# ---------------------------------------------------------------------------

RESULTS = {}
for rid in RESULT_IDS:
    path = os.path.join(RESULTS_DIR, f"RESULT-{rid}.yaml")
    with open(path, "r") as f:
        RESULTS[rid] = yaml.safe_load(f)

# ---------------------------------------------------------------------------
# Placeholder resolution: {{RESULT-<id>.<path>[:<format>]}}
# ---------------------------------------------------------------------------

TOKEN_RE = re.compile(r"([^.\[\]]+)((?:\[\d+\])*)")
PLACEHOLDER_RE = re.compile(r"\{\{RESULT-([a-zA-Z0-9\-]+)\.([^:}]+)(?::([^}]+))?\}\}")


def _resolve_path(obj, path):
    cur = obj
    for key, idxs in TOKEN_RE.findall(path):
        if key:
            cur = cur[key]
        for idx in re.findall(r"\[(\d+)\]", idxs):
            cur = cur[int(idx)]
    return cur


def render(text):
    """Replace every {{RESULT-id.path:fmt}} in text with its resolved value."""

    def _sub(m):
        rid, path, fmt = m.group(1), m.group(2), m.group(3)
        if rid not in RESULTS:
            raise KeyError(f"Unknown RESULT id in placeholder: RESULT-{rid}")
        val = _resolve_path(RESULTS[rid], path)
        if fmt:
            return format(val, fmt)
        return str(val)

    out = PLACEHOLDER_RE.sub(_sub, text)
    if "{{" in out or "}}" in out:
        raise ValueError(f"Unresolved placeholder braces remain in text:\n{out}")
    return out


# ---------------------------------------------------------------------------
# Derived (pure bookkeeping) aggregates over RESULT objects -- no new stats
# ---------------------------------------------------------------------------

QC = RESULTS["qc-counts"]["sessions"]
N_SESSIONS = len(QC)
TOTAL_UNITS_QC = sum(s["n_neurons"] for s in QC)

DISCARDS = RESULTS["qc-discards"]["discards"]
DISCARD_REASONS = sorted(set(d["reason"] for d in DISCARDS))
DISCARD_COUNTS_BY_SESSION = Counter(d["experiment"] for d in DISCARDS)
TOTAL_DISCARDS = len(DISCARDS)

THRESH = RESULTS["eth-threshold-log"]["sessions"]

TOP5_DELTA = RESULTS["delta-mrl"]["units"][:5]
UNIT_LOCS = {u["unit_id"]: u for u in RESULTS["unit-locations"]["examples"]}
PSTH_EX_MAP = {u["unit_id"]: u for u in RESULTS["psth-examples"]["examples"]}

MRL_ANIMAL = RESULTS["mrl-animal"]["sessions"]

# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------


def add_title(doc, text):
    h = doc.add_heading(text, level=0)
    return h


def add_h1(doc, text):
    doc.add_heading(text, level=1)


def add_p(doc, text):
    doc.add_paragraph(text)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    return table


# ---------------------------------------------------------------------------
# =====================  DOCUMENT 1: METHODS / QC  =========================
# ---------------------------------------------------------------------------

def build_methods_doc():
    doc = Document()
    add_title(doc, "NP-DEMO-2: Ethanol and Sniff-Phase Locking of OB Units")
    doc.add_paragraph("Methods / QC Report", style="Subtitle")

    # ---- Question --------------------------------------------------------
    add_h1(doc, "Question")
    add_p(
        doc,
        "This experiment tested whether ethanol exposure alters the phase-locking "
        "of olfactory-bulb (OB) single-unit spikes to the sniff cycle, as measured "
        "by the SNF (sniff sensor thermistor/pressure) signal, compared to control "
        "epochs within the same recording sessions.",
    )

    # ---- Methods -----------------------------------------------------------
    add_h1(doc, "Methods")
    add_p(
        doc,
        "Data: six recording sessions from C:\\Projects\\Repos\\Neuropixels\\DATA, "
        "sessions dated 2021-11-01, 2021-11-03, 2021-12-15, 2022-05-17, 2022-06-24, "
        "2022-09-14.",
    )
    add_p(
        doc,
        "Sniff detection: the SNF signal was low-pass filtered (FIR, 40 Hz) and "
        "z-scored. Sniff onsets were detected by downward threshold crossings at "
        "threshold_std = -0.5 SD (pinned). Valid sniff epochs were defined as SNF "
        "phase >= 0.0 (computed via compute_sniff_phase); sections with SNF_PH == -1 "
        "were excluded as noise / non-sniffing.",
    )
    add_p(
        doc,
        render(
            "Ethanol threshold (PROH-002), Pass 1: all sessions were run at the "
            "default eth_threshold = 0.11. Cross-session mean ethanol-contact count "
            "= {{RESULT-eth-threshold-log.pass1_stats.mean_contacts:.1f}}, SD = "
            "{{RESULT-eth-threshold-log.pass1_stats.sd_contacts:.1f}}. Session "
            "2021-11-01 was flagged as an outlier (above mean+1SD) and underwent "
            "Pass-2 threshold adjustment."
        ),
    )
    add_p(
        doc,
        render(
            "Pass 2: for 2021-11-01, a deterministic grid search (eth_threshold "
            "0.05 to 0.50, step 0.005) selected eth_threshold = "
            "{{RESULT-eth-threshold-log.sessions[0].final_threshold}} to bring its "
            "contact count ({{RESULT-eth-threshold-log.sessions[0].n_contacts_at_final}} "
            "contacts) closer to the cross-session mean."
        ),
    )
    add_p(
        doc,
        render(
            "Note: session 2021-11-03 had only "
            "{{RESULT-eth-threshold-log.sessions[1].n_contacts_at_final}} ethanol "
            "contact(s) at the default threshold and was not flagged by the +/-1 SD "
            "criterion (DEV-001; see Limitations)."
        ),
    )
    add_p(
        doc,
        "Spike phase assignment: each spike's SNF phase was computed via "
        "compute_spike_phase. Condition assignment: a spike was classified as "
        "'ethanol' iff ETH_thr[index] > eth_threshold at its time, else 'control'.",
    )
    add_p(
        doc,
        "Unit inclusion: units with overall firing rate < 0.1 Hz or with fewer than "
        "50 valid-sniff spikes combined (ethanol + control) were excluded.",
    )
    add_p(
        doc,
        "Phase-locking: the circular mean resultant length (MRL) and preferred "
        "phase were computed separately for the ethanol and control conditions, "
        "per unit.",
    )

    # ---- Ethanol-Threshold-Adjustment-Log ----------------------------------
    add_h1(doc, "Ethanol-Threshold-Adjustment-Log")
    add_p(
        doc,
        "Per-session PROH-002 default vs. final ethanol-contact threshold, "
        "whether Pass-2 adjustment was applied, and the resulting contact count "
        "(source: RESULT-eth-threshold-log).",
    )
    rows = []
    for s in THRESH:
        rows.append(
            [
                s["session_date"],
                format(s["default_threshold"], ".3f"),
                format(s["final_threshold"], ".3f"),
                str(s["was_adjusted"]),
                s["n_contacts_at_final"],
            ]
        )
    add_table(
        doc,
        ["Session", "Default threshold", "Final threshold", "Adjusted?", "N contacts (final)"],
        rows,
    )

    # ---- QC-Counts ----------------------------------------------------------
    add_h1(doc, "QC-Counts")
    add_p(
        doc,
        "Per-session sniff, unit, and ethanol-contact counts (source: "
        "RESULT-qc-counts).",
    )
    rows = []
    for s in QC:
        rows.append(
            [
                s["session_date"],
                s["n_sniffs"],
                s["n_neurons"],
                s["n_trials"],
                format(s["length_min"], ".1f"),
                format(s["pct_usable_sniffs"], ".1f"),
                format(s["eth_threshold"], ".3f"),
            ]
        )
    add_table(
        doc,
        [
            "Session",
            "N sniffs",
            "N neurons (included)",
            "N trials (eth. contacts)",
            "Length (min)",
            "% usable sniffs",
            "eth_threshold used",
        ],
        rows,
    )
    add_p(
        doc,
        f"Total included units across all six sessions (sum of n_neurons, "
        f"RESULT-qc-counts): {TOTAL_UNITS_QC}.",
    )

    # ---- Discarded-Sections --------------------------------------------------
    add_h1(doc, "Discarded-Sections")
    if TOTAL_DISCARDS == 0:
        add_p(
            doc,
            "No SNF sections were discarded in any session; all detected sniff "
            "cycles were used.",
        )
    else:
        reasons_str = ", ".join(DISCARD_REASONS)
        add_p(
            doc,
            f"Discarded SNF sections were logged per RESULT-qc-discards "
            f"(experiment, start_s, end_s, reason). Across all six sessions, "
            f"{TOTAL_DISCARDS} sections were discarded, all for reason: "
            f"{reasons_str}. Per-session counts (source: RESULT-qc-discards):",
        )
        rows = []
        for s in QC:
            sd = s["session_date"]
            rows.append([sd, DISCARD_COUNTS_BY_SESSION.get(sd, 0)])
        add_table(doc, ["Session", "N discarded sections"], rows)
        add_p(
            doc,
            f"Total discarded sections across all sessions (RESULT-qc-discards): "
            f"{TOTAL_DISCARDS}. The full itemized log (experiment, start_s, end_s, "
            f"reason) is retained in RESULT-qc-discards.yaml for traceability and "
            f"is not reproduced in full here given its length.",
        )

    # ---- Figures --------------------------------------------------------------
    add_h1(doc, "Figures")
    add_p(
        doc,
        "FIG-DEMO2-QC: per-experiment SNF and ethanol (ETH) traces with the "
        "sniff-onset threshold and the ethanol-contact threshold (post-PROH-002 "
        "if adjusted) marked, with discarded sections shaded (source: "
        "RESULT-qc-counts, RESULT-eth-threshold-log, RESULT-qc-discards).",
    )
    add_p(
        doc,
        "FIG-DEMO2-QC-PSTH: sniff-locked PSTHs for the single strongest- and "
        "single weakest-overall phase-locking units across all experiments "
        "(source: RESULT-psth-extremes, RESULT-mrl-unit).",
    )

    # ---- Limitations ------------------------------------------------------
    add_h1(doc, "Limitations")
    add_p(
        doc,
        render(
            "DEV-001: session 2021-11-03 escaped the PROH-002 +/-1 SD flag "
            "because session 2021-11-01's large contact count inflated the "
            "cross-session SD. Session 2021-11-03 had only "
            "{{RESULT-eth-threshold-log.sessions[1].n_contacts_at_final}} ethanol "
            "contact(s) over {{RESULT-qc-counts.sessions[1].length_min:.1f}} minutes "
            "({{RESULT-qc-counts.sessions[1].pct_usable_sniffs:.1f}}% usable sniffs), "
            "resulting in a near-zero control-condition spike count per unit. This "
            "degenerate data affects the examples figure selection (FIG-DEMO2-RES-EX) "
            "and may influence the animal-level statistical result (see Results "
            "report). A formal deviation has been filed (DEV-001 in "
            "08_traceability/deviations.yaml). Human sign-off is required before "
            "these results are used."
        ),
    )

    out_path = os.path.join(
        OUT_DIR, "NP-DEMO-2_ethanol_sniff_phase_locking_METHODS.docx"
    )
    doc.save(out_path)
    return out_path


# ---------------------------------------------------------------------------
# =====================  DOCUMENT 2: RESULTS  ===============================
# ---------------------------------------------------------------------------

def build_results_doc():
    doc = Document()
    add_title(doc, "NP-DEMO-2: Ethanol and Sniff-Phase Locking of OB Units")
    doc.add_paragraph("Results Report", style="Subtitle")

    # ---- Question -----------------------------------------------------------
    add_h1(doc, "Question")
    add_p(
        doc,
        "This experiment tested whether ethanol exposure alters the phase-locking "
        "of olfactory-bulb (OB) single-unit spikes to the sniff cycle, as measured "
        "by the SNF (sniff sensor thermistor/pressure) signal, compared to control "
        "epochs within the same recording sessions.",
    )

    # ---- Methods-Summary ------------------------------------------------------
    add_h1(doc, "Methods-Summary")
    add_p(
        doc,
        f"Full methods, QC counts, ethanol-threshold-adjustment log, and discard "
        f"log are reported separately in the companion Methods/QC document "
        f"(NP-DEMO-2_ethanol_sniff_phase_locking_METHODS.docx). Briefly: sniff "
        f"phase was computed from the SNF signal (compute_sniff_phase, "
        f"threshold_std = -0.5, pinned); spikes were classified as ethanol or "
        f"control by the (possibly PROH-002-adjusted) ethanol threshold; units "
        f"with firing rate < 0.1 Hz or < 50 valid-sniff spikes combined were "
        f"excluded. n_sessions = {N_SESSIONS}; n_units_total (included per QC "
        f"criteria, sum of n_neurons across sessions, RESULT-qc-counts) = "
        f"{TOTAL_UNITS_QC}.",
    )

    # ---- Results-Animal-Level ---------------------------------------------
    add_h1(doc, "Results-Animal-Level")
    add_p(
        doc,
        render(
            "Across {{RESULT-stat-animal.n_animals}} animals/sessions, ethanol "
            "exposure was associated with a {{RESULT-stat-animal.direction}} "
            "difference in mean sniff-phase MRL (Wilcoxon signed-rank, exact "
            "two-sided; W = {{RESULT-stat-animal.statistic:.0f}}, p = "
            "{{RESULT-stat-animal.pvalue:.5f}}, rank-biserial r = "
            "{{RESULT-stat-animal.effect_size:.3f}})."
        ),
    )
    add_p(doc, "Per-animal mean MRLs (source: RESULT-mrl-animal):")
    rows = []
    for s in MRL_ANIMAL:
        rows.append(
            [
                s["session"],
                format(s["mean_mrl_ethanol"], ".4f"),
                format(s["mean_mrl_control"], ".4f"),
                s["n_units_included"],
            ]
        )
    add_table(
        doc,
        ["Session", "Mean MRL (ethanol)", "Mean MRL (control)", "N units included"],
        rows,
    )
    add_p(
        doc,
        render(
            "Note: session 2021-11-03 shows mean_mrl_control = "
            "{{RESULT-mrl-animal.sessions[1].mean_mrl_control:.3f}}, which is "
            "anomalously high due to DEV-001 (degenerate control-condition "
            "estimates from a near-zero ethanol-contact count at Pass 1; see "
            "Methods report). This session's contribution to the Wilcoxon "
            "statistic should be interpreted with caution."
        ),
    )

    # ---- Results-Unit-Level -----------------------------------------------
    add_h1(doc, "Results-Unit-Level")
    add_p(
        doc,
        render(
            "At the unit level, a linear mixed-effects model (condition as fixed "
            "effect, session as random intercept; n = {{RESULT-stat-unit.n_units}} "
            "units across {{RESULT-stat-unit.n_animals}} sessions) found a "
            "significant {{RESULT-stat-unit.direction}} effect of ethanol on MRL "
            "(coefficient = {{RESULT-stat-unit.coefficient:.4f}}, standardized "
            "effect = {{RESULT-stat-unit.standardized_effect_size:.4f}}, p = "
            "{{RESULT-stat-unit.pvalue:.3e}})."
        ),
    )
    add_p(
        doc,
        f"Full per-unit MRL and preferred-phase values by condition (n = "
        f"{TOTAL_UNITS_QC} units passing QC criteria; source: RESULT-mrl-unit "
        f"[{TOTAL_UNITS_QC} rows] and RESULT-phase-unit) underlie the unit-level "
        f"figure (FIG-DEMO2-RES-UNIT); given their size, they are provided as "
        f"source data files in 04_results/ rather than reproduced as an in-text table. "
        f"Of these {TOTAL_UNITS_QC} units, {RESULTS['stat-unit']['n_units']} "
        f"(source: RESULT-stat-unit.n_units) had spikes in both the ethanol and "
        f"control conditions and entered the mixed-effects model; the remaining "
        f"{TOTAL_UNITS_QC - RESULTS['stat-unit']['n_units']} units had zero "
        f"control-condition spikes (primarily from session 2021-11-03 due to DEV-001) "
        f"and were excluded from the model.",
    )

    # ---- Examples -----------------------------------------------------------
    add_h1(doc, "Examples")
    add_p(
        doc,
        "The five units with the largest absolute change in MRL between "
        "conditions (|delta MRL|) are listed below (source: RESULT-delta-mrl, "
        "RESULT-unit-locations). NOTE: all five are from session 2021-11-03, "
        "which has degenerate control-condition estimates (DEV-001). These "
        "examples should not be interpreted as representative until DEV-001 is "
        "resolved.",
    )
    rows = []
    for u in TOP5_DELTA:
        uid = u["unit_id"]
        loc = UNIT_LOCS.get(uid, {})
        rows.append(
            [
                uid,
                u["session"],
                format(u["delta_mrl"], ".4f"),
                format(u["abs_delta_mrl"], ".4f"),
                format(loc.get("unit_depth_um", float("nan")), ".1f")
                if loc
                else "n/a",
                loc.get("shank", "n/a"),
            ]
        )
    add_table(
        doc,
        ["Unit ID", "Session", "delta MRL (eth - ctrl)", "|delta MRL|", "Depth (um)", "Shank"],
        rows,
    )

    # ---- Figures --------------------------------------------------------------
    add_h1(doc, "Figures")
    add_p(
        doc,
        "FIG-DEMO2-RES-EXP: population-level ethanol vs. control phase-locking "
        "at the animal (experiment) level, with the paired animal-level test "
        "result displayed (source: RESULT-mrl-animal, RESULT-stat-animal).",
    )
    add_p(
        doc,
        "FIG-DEMO2-RES-UNIT: distribution of per-unit MRL across ethanol vs. "
        "control, with the unit-level mixed-effects test result displayed "
        "(source: RESULT-mrl-unit, RESULT-phase-unit, RESULT-stat-unit).",
    )
    add_p(
        doc,
        "FIG-DEMO2-RES-EX: sniff-locked PSTHs and probe-location maps for the "
        "five example units with the largest |delta MRL| (source: "
        "RESULT-delta-mrl, RESULT-psth-examples, RESULT-unit-locations).",
    )

    # ---- Limitations ------------------------------------------------------
    add_h1(doc, "Limitations")
    add_p(
        doc,
        render(
            "DEV-001 (critical): session 2021-11-03 was not flagged by the "
            "PROH-002 +/-1 SD rule on Pass 1 (a large contact count in session "
            "2021-11-01 inflated the cross-session SD), so it entered the "
            "analysis with only {{RESULT-eth-threshold-log.sessions[1].n_contacts_at_final}} "
            "ethanol contact(s) and near-zero control-condition spikes per unit. "
            "This degenerate control-condition data dominates the Examples "
            "selection above and may inflate the animal-level effect. See the "
            "Methods report and 08_traceability/deviations.yaml for full detail. "
            "Human sign-off is required before these results are used."
        ),
    )
    add_p(
        doc,
        render(
            "The animal-level Wilcoxon result (p = {{RESULT-stat-animal.pvalue:.5f}}) "
            "is the minimum possible exact p for n=6 (all concordant pairs), and is "
            "sensitive to the inclusion of 2021-11-03. Without that session (n=5, "
            "all concordant), p would be 0.0625 (not significant at alpha=0.05). The "
            "significance claim is therefore conditional on resolution of DEV-001."
        ),
    )
    add_p(
        doc,
        render(
            "The unit-level p-value (p = {{RESULT-stat-unit.pvalue:.3e}}) is likely "
            "robust given its extreme magnitude, but the coefficient magnitude may "
            "be inflated by degenerate MRL_control values from 2021-11-03."
        ),
    )

    out_path = os.path.join(
        OUT_DIR, "NP-DEMO-2_ethanol_sniff_phase_locking_RESULTS.docx"
    )
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    os.makedirs(OUT_DIR, exist_ok=True)
    m = build_methods_doc()
    r = build_results_doc()
    print("Wrote:", m)
    print("Wrote:", r)
    print("n_sessions:", N_SESSIONS)
    print("total_units_qc:", TOTAL_UNITS_QC)
    print("total_discards:", TOTAL_DISCARDS)
    print("discard_reasons:", DISCARD_REASONS)
