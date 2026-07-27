"""
render_contract_docx.py
Produce _pipeline/NP-DEMO-3/02_contract/contract_v001.docx from contract_v001.yaml.
Also updates manifest.yaml with the docx SHA256.
"""
from __future__ import annotations
import hashlib, pathlib
import yaml
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT     = pathlib.Path(__file__).parent
CONTRACT = ROOT / "02_contract" / "contract_v001.yaml"
OUT_DOCX = ROOT / "02_contract" / "contract_v001.docx"
MANIFEST = ROOT / "02_contract" / "manifest.yaml"

data = yaml.safe_load(CONTRACT.read_text(encoding="utf-8"))
ana  = data["analysis"]

doc = Document()

# ── page margins ────────────────────────────────────────────────────────────
sec = doc.sections[0]
sec.top_margin    = Inches(1.0)
sec.bottom_margin = Inches(1.0)
sec.left_margin   = Inches(1.25)
sec.right_margin  = Inches(1.25)


# ── helper functions ─────────────────────────────────────────────────────────
def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def h3(text):
    doc.add_heading(text, level=3)

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(str(text))
    run.bold   = bold
    run.italic = italic
    return p

def bullet(text):
    doc.add_paragraph(str(text), style="List Bullet")

def kv(key, value, indent=False):
    """Key: Value paragraph."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    r1 = p.add_run(f"{key}: ")
    r1.bold = True
    p.add_run(str(value) if value is not None else "—")
    return p

def hr():
    """Thin horizontal rule via bottom border on an empty paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)

def simple_table(rows, header=None):
    cols = len(rows[0]) if rows else (len(header) if header else 1)
    tbl = doc.add_table(rows=0, cols=cols)
    tbl.style = "Table Grid"
    if header:
        row = tbl.add_row()
        for i, h in enumerate(header):
            cell = row.cells[i]
            cell.text = str(h)
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    for r in rows:
        row = tbl.add_row()
        for i, val in enumerate(r):
            cell = row.cells[i]
            cell.text = str(val) if val is not None else "—"
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    doc.add_paragraph()   # spacer after table


# ════════════════════════════════════════════════════════════════════════════
# TITLE
# ════════════════════════════════════════════════════════════════════════════
title_p = doc.add_heading("Analysis Contract", level=0)
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run(f"Task: {data['task_id']}   |   Version: {data['contract_version']}   |   Status: DRAFT — PENDING APPROVAL")
r.bold = True
r.font.size = Pt(12)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 1 — OVERVIEW
# ════════════════════════════════════════════════════════════════════════════
h1("1. Overview")
kv("Task ID",            data["task_id"])
kv("Contract version",   data["contract_version"])
kv("Risk tier",          ana["risk"].upper())
kv("Confirmatory status", ana["status"])
kv("Randomization seed", ana["randomization_seed"])
doc.add_paragraph()

h2("Scientific Question")
para(ana["scientific_question"])
doc.add_paragraph()

h2("Experimental Structure")
para(ana["experimental_structure"])
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 2 — DATA PROVENANCE
# ════════════════════════════════════════════════════════════════════════════
h1("2. Data Provenance")
dp = ana["data_provenance"]
kv("Root directory", dp["root"])
kv("Sessions", ", ".join(dp["sessions"]))
kv("Source refs", ", ".join(dp.get("source_refs", [])))
kv("Pinning status", dp.get("pinning_status", ""))
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 3 — INPUTS & FUNCTIONS
# ════════════════════════════════════════════════════════════════════════════
h1("3. Inputs and Required Functions")

h2("Inputs")
for inp in ana["inputs"]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(inp.get("object", inp.get("from", "?")))
    r.bold = True
    desc = inp.get("description") or inp.get("from") or ""
    if inp.get("fields"):
        desc += f" — fields: {', '.join(inp['fields'])}"
    if inp.get("units"):
        desc += f" (units: {inp['units']})"
    if inp.get("rate_Hz_source"):
        desc += f", rate: {inp['rate_Hz_source']}"
    if inp.get("source"):
        desc += f" — {inp['source']}"
    p.add_run(f": {desc}")
doc.add_paragraph()

h2("Validated Kernel Functions (reuse; do NOT edit)")
for fn in ana["required_validated_functions"]:
    bullet(fn)
doc.add_paragraph()

h2("New Kernel Functions Required")
for fn in ana.get("new_kernel_functions_required", []):
    h3(fn["name"])
    kv("Proposed signature", fn["proposed_signature"])
    kv("Algorithm", fn["algorithm"])
    kv("Fixture approach", fn["fixture_approach"])
    kv("Constraint", fn["constraint"])
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 4 — ANALYSIS DEFINITION
# ════════════════════════════════════════════════════════════════════════════
h1("4. Analysis Definition")
para(ana["analysis_definition"])
doc.add_paragraph()

h2("Preprocessing Steps")
for i, step in enumerate(ana["preprocessing"], 1):
    p = doc.add_paragraph(style="List Number")
    p.add_run(step)
doc.add_paragraph()

h2("Aggregation")
for step in ana["aggregation"]:
    bullet(step)
doc.add_paragraph()

h2("Statistical Model")
para(ana["statistical_model"])
doc.add_paragraph()

h2("Multiple Comparisons")
para(ana["multiple_comparison"])
doc.add_paragraph()

h2("Assumptions")
for a in ana["assumptions"]:
    bullet(a)
doc.add_paragraph()

h2("Inclusion / Exclusion Criteria")
para(ana["inclusion_exclusion"])
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 5 — REQUIRED OUTPUTS & ACCEPTANCE CRITERIA
# ════════════════════════════════════════════════════════════════════════════
h1("5. Required Outputs and Acceptance Criteria")

# Build a combined table
ac_map = {ac["output_id"]: ac for ac in ana["acceptance_criteria"]}
rows = []
for ro in ana["required_outputs"]:
    rid = ro["result_id"]
    ac  = ac_map.get(rid, {})
    rows.append([
        rid,
        ro["type"],
        ro["desc"][:120] + ("…" if len(ro["desc"]) > 120 else ""),
        f"L{ac.get('level','?')}",
        ac.get("comparison_method", ""),
        ac.get("tolerance", "")[:80] + ("…" if len(ac.get("tolerance","")) > 80 else ""),
    ])

simple_table(
    rows,
    header=["Result ID", "Type", "Description", "Level", "Method", "Tolerance"],
)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 6 — PROHIBITED CHANGES & FAILURE CONDITIONS
# ════════════════════════════════════════════════════════════════════════════
h1("6. Prohibited Changes")
for pc in ana["prohibited_changes"]:
    bullet(pc)
doc.add_paragraph()

h1("7. Failure Conditions")
for fc in ana["failure_conditions"]:
    bullet(fc)
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 7 — FIGURES
# ════════════════════════════════════════════════════════════════════════════
h1("8. Figures")

for fig in data["figures"]:
    h2(fig["figure_id"])
    kv("Title",       fig["title"])
    kv("Deliverable", fig.get("deliverable", ""))
    kv("Required",    fig["required"])
    kv("Purpose",     fig["scientific_purpose"])
    kv("Source results", ", ".join(fig.get("source_results", [])))

    ax = fig.get("axis_requirements", {})
    if ax:
        kv("X label",   ax.get("x_label", ""))
        kv("Y label",   ax.get("y_label", ""))
        kv("Truncation permitted", ax.get("truncation_permitted", False))

    h3("Caption requirements")
    for cr in fig.get("caption_requirements", []):
        bullet(cr)

    sk = fig.get("sketch_guidance")
    if sk:
        p = doc.add_paragraph()
        r = p.add_run("Sketch guidance (NON-BINDING): ")
        r.bold = True
        r.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
        p.add_run(sk.get("intent", ""))

    q = fig.get("quality", {})
    kv("Quality", f"vector={q.get('vector')}, dpi={q.get('raster_dpi')}, colorblind={q.get('colorblind_review')}, min_font={q.get('min_font')}pt")
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 8 — REPORTS
# ════════════════════════════════════════════════════════════════════════════
h1("9. Report Deliverables")

for rpt in data.get("reports", []):
    h2(f"{rpt.get('role','report').upper()} — {rpt.get('output_filename','')}")
    kv("Audience",        rpt.get("intended_audience", ""))
    kv("Scientific status", rpt.get("scientific_status", ""))
    if rpt.get("separate_from_results_report"):
        para("This is a SEPARATE document from the results report.", bold=True)

    h3("Required sections")
    for s in rpt.get("required_sections", []):
        bullet(s)

    h3("Required figures")
    for f in rpt.get("required_figures", []):
        bullet(f)

    h3("Required tables")
    for t in rpt.get("required_tables", []):
        bullet(t)

    rr = rpt.get("reporting_rules", {})
    h3("Reporting rules")
    for k, v in rr.items():
        kv(k, v, indent=True)

    for ap in rpt.get("appendices", []):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run("Appendix: ").bold = True
        p.add_run(ap)
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 9 — OPEN DECISIONS
# ════════════════════════════════════════════════════════════════════════════
h1("10. Open Decisions")
od = data.get("open_decisions", [])
if not od:
    para("None — all open decisions resolved. Contract is eligible for gate-A approval.", bold=True)
else:
    for item in od:
        h2(item.get("id", "?"))
        para(item.get("question", ""), italic=True)
        if item.get("resolution"):
            kv("Resolution", item["resolution"])
        kv("Blocks approval", item.get("blocks_approval", False))
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 10 — APPROVAL BLOCK
# ════════════════════════════════════════════════════════════════════════════
hr()
h1("11. Approval")
para(
    "By signing below the approver confirms that this contract accurately captures "
    "the scientific question, methods, outputs, figures, and reports, and that all "
    "open decisions are resolved. Approval hash-locks this version; no substantive "
    "change may be made without issuing a new version (v002, v003, …).",
    italic=True,
)
doc.add_paragraph()

for label in ["Approver name", "Date", "Signature"]:
    p = doc.add_paragraph()
    p.add_run(f"{label}:  ").bold = True
    p.add_run("_" * 50)

doc.add_paragraph()
kv("Contract YAML SHA256", "40c6fda4412e5995dcce5e4e374ea3e4e27e2b920b2d31dfb3fabde31bff218e")
kv("Source request SHA256", "86e3cb81292f213b1f874be052dc56481861c669021fe561a4f602e13fa94aa7")

# ════════════════════════════════════════════════════════════════════════════
# SAVE & UPDATE MANIFEST
# ════════════════════════════════════════════════════════════════════════════
doc.save(OUT_DOCX)
docx_hash = hashlib.sha256(OUT_DOCX.read_bytes()).hexdigest()

manifest = yaml.safe_load(MANIFEST.read_text(encoding="utf-8"))
manifest["docx_sha256"] = docx_hash
MANIFEST.write_text(yaml.dump(manifest, sort_keys=True, allow_unicode=True), encoding="utf-8")

print(f"contract_v001.docx written: {OUT_DOCX}")
print(f"docx_sha256: {docx_hash}")
print("manifest.yaml updated.")
