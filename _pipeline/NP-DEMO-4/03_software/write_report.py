"""
NP-DEMO-4 report generator.

Generates NP-DEMO-4_methods.docx and NP-DEMO-4_results.docx from result-object
YAML files under 04_results/. All numeric values are read from those result
objects at render time -- no statistic is typed by hand in this script.

Run with:
  C:/Projects/Repos/Agentic Research/Research Setup/research_workflow/vras/Scripts/python.exe write_report.py
"""

import os
import yaml
from docx import Document
from docx.shared import Pt

REPO_ROOT   = r"C:\Projects\Repos\Neuropixels\_pipeline\NP-DEMO-4"
RESULTS_DIR = REPO_ROOT + r"\04_results"
REPORT_DIR  = REPO_ROOT + r"\07_report"
FIG_DIR     = REPO_ROOT + r"\06_figures\figure_data"

os.makedirs(REPORT_DIR, exist_ok=True)


def load_yaml(path):
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return yaml.safe_load(f)


# ---------------------------------------------------------------------------
# Load all result objects (source of every injected number in this report)
# ---------------------------------------------------------------------------
sniff_stat      = load_yaml(RESULTS_DIR + r"\RESULT-sniff-stat.yaml")
unit_inclusion  = load_yaml(RESULTS_DIR + r"\RESULT-unit-inclusion.yaml")
methods_table   = load_yaml(RESULTS_DIR + r"\RESULT-methods-table.yaml")
eth_mask        = load_yaml(RESULTS_DIR + r"\RESULT-eth-mask.yaml")
figure_manifest = load_yaml(FIG_DIR + r"\figure_data_manifest.yaml")

# Convenience: session list in order (matches RESULT-methods-table and
# per-animal arrays in RESULT-sniff-stat, which are aligned by session)
SESSIONS = [row["session"] for row in methods_table]
N_ANIMALS = int(sniff_stat["n_animals"])       # 6
W_STAT    = float(sniff_stat["W"])              # 7.0
P_EXACT   = float(sniff_stat["p_exact"])        # 0.5625
RBR       = float(sniff_stat["rank_biserial_r"])  # 0.3333...
SIG       = bool(sniff_stat["significant"])     # False
CONTACT_MEANS = sniff_stat["contact_means"]     # list of 6 floats
CONTROL_MEANS = sniff_stat["control_means"]     # list of 6 floats

# Unit inclusion: build lookup {session: {n_rec, n_inc}}
ui_lookup = {row["session"]: row for row in unit_inclusion}

# Figure manifest caption injections
N_TRIALS_FIG = int(figure_manifest["figures"]["FIG_DEMO4_1_SNIFF"]["n_total_trials"])
N_SESS_FIG   = int(figure_manifest["figures"]["FIG_DEMO4_1_SNIFF"]["n_sessions"])


def set_base_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def add_title(doc, text):
    return doc.add_heading(text, level=0)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text):
    return doc.add_paragraph(text)


# ===========================================================================
# NP-DEMO-4_methods.docx
# ===========================================================================
doc_m = Document()
set_base_style(doc_m)

add_title(doc_m, "NP-DEMO-4: Methods")

add_heading(doc_m, "Research Question", level=1)
add_para(
    doc_m,
    "This analysis addresses whether instantaneous sniff rate during the first "
    "ethanol contact on a trial differs from sniff rate during a pre-valve "
    "control period (trial timestamp, TS, 0-10 s). Related descriptive outputs "
    "characterize how sniff rate, olfactory bulb population firing rate, and "
    "the ethanol sensor signal evolve across the full trial window.",
)

add_heading(doc_m, "Subjects and Sessions", level=1)
session_date_str = ", ".join(SESSIONS)
add_para(
    doc_m,
    f"Data comprise {N_ANIMALS} Neuropixels olfactory-bulb (OB) recording "
    f"sessions, one animal per session, recorded on the following dates: "
    f"{session_date_str}.",
)

add_heading(doc_m, "Data Acquisition", level=1)
add_para(
    doc_m,
    "Neuropixels recordings were made from olfactory bulb (OB). Behavior was "
    "controlled and timestamped by a LabView acquisition system, which recorded "
    "the following synchronized signals: SNF (raw sniff-sensor pressure trace, "
    "~125 Hz), ETH (ethanol sensor trace, normalized 0-1), TR (LabView trial "
    "register, integer trial index per sample), and TS (within-trial timestamp "
    "in milliseconds, TS = 0 at trial start).",
)

add_heading(doc_m, "Trial Selection and Segmentation", level=1)
add_para(
    doc_m,
    "Only odd TR values were treated as valid LabView trials; even TR values "
    "(inter-trial intervals) were excluded from all analyses and figures. "
    "Contiguous samples sharing the same odd TR value were segmented into a "
    "trial spanning TS = 0-40 s, with the first (corrupted) TS sample of each "
    "trial corrected to 0. Duplicate TS samples arising from a data-acquisition "
    "artifact were identified and removed prior to trial segmentation (see "
    "Deviations, below).",
)

add_heading(doc_m, "Sniff Rate Computation", level=1)
add_para(
    doc_m,
    "Instantaneous sniff rate was computed from the raw SNF trace exclusively "
    "(never from LFP, per PROH-001). The SNF trace was passed through a 21-tap "
    "FIR lowpass filter (40 Hz cutoff, Hamming window) applied with zero-phase "
    "filtering (filtfilt), then z-scored (ddof = 1). Sniff onsets were detected "
    "as downward zero-crossings of the z-scored trace at a threshold of -0.5, "
    "with a minimum inter-sniff-interval (MIN_ISI) rejection criterion applied "
    "to exclude spuriously close onsets. Instantaneous sniff rate at each "
    "LabView sample was obtained by interpolating 1/ISI (inverse inter-onset "
    "interval) to the LabView sample times.",
)

add_heading(doc_m, "Unit Inclusion Criteria", level=1)
ui_str_parts = []
for sess in SESSIONS:
    row = ui_lookup[sess]
    ui_str_parts.append(
        f"{sess} — {row['n_units_recorded']} recorded, "
        f"{row['n_units_included']} included"
    )
add_para(
    doc_m,
    "OB units were included in firing-rate analyses only if they met both of "
    "the following criteria simultaneously, computed over the entire session: "
    "overall mean firing rate >= 0.1 Hz, and total spike count >= 5000. "
    "Per-session unit counts (units recorded and units meeting inclusion "
    "criteria) are as follows: "
    + "; ".join(ui_str_parts)
    + ".",
)

add_heading(doc_m, "Ethanol Contact Detection", level=1)
add_para(
    doc_m,
    "Ethanol contact events were detected from the ETH sensor trace after "
    "subtracting the mean of the entire session-wide ETH trace, thresholded "
    "strictly greater than 0.05 (fixed rule, applied identically to all "
    "sessions, no per-experiment adjustment). A discrete contact event was "
    "defined as a contiguous run of samples above this threshold. For the "
    "sniff-rate comparison (CON-003, below), the first discrete ETH contact "
    "event within each LabView trial was used.",
)

add_heading(doc_m, "Statistical Test (CON-003)", level=1)
add_para(
    doc_m,
    f"For each LabView trial, mean instantaneous sniff rate was computed within "
    f"(a) the first discrete ETH contact period and (b) the pre-valve control "
    f"window (TS = 0-10 s, excluding any samples overlapping a contact event). "
    f"These per-trial means were averaged within each session/animal to yield "
    f"one contact-period mean and one control-period mean per animal "
    f"(n = {N_ANIMALS} animals). The two-sided exact Wilcoxon signed-rank test "
    f"(scipy.stats.wilcoxon, method='exact', zero_method='wilcox') was applied "
    f"to the {N_ANIMALS} paired per-animal means. Matched-pairs rank-biserial r "
    f"was computed as r = (W+ - W-)/(W+ + W-) as the effect size. The "
    f"significance threshold was set at alpha = 0.05.",
)

add_heading(doc_m, "Deviations", level=1)
add_para(
    doc_m,
    "DEV-001: during implementation, TS sample duplication was detected within "
    "the trial-segmentation procedure (repeated TS values arising from a "
    "startup artifact in one session), which caused per-trial sample counts to "
    "be inflated. This was corrected by detecting TS resets within each odd-TR "
    "block and removing duplicate TS samples prior to extracting the 0-40 s "
    "trial window. This deviation was reviewed and resolved prior to analysis "
    "finalization; all reported results reflect the corrected trial "
    "segmentation.",
)

doc_m.save(REPORT_DIR + r"\NP-DEMO-4_methods.docx")
print(f"Wrote: {REPORT_DIR}\\NP-DEMO-4_methods.docx")


# ===========================================================================
# NP-DEMO-4_results.docx
# ===========================================================================
doc_r = Document()
set_base_style(doc_r)

add_title(doc_r, "NP-DEMO-4: Results")

# ---------------------------------------------------------------------------
# 1. Per-session summary table
# ---------------------------------------------------------------------------
add_heading(doc_r, "Per-Session Summary", level=1)
add_para(
    doc_r,
    "Table 1 summarizes, per session, the number of LabView trials analyzed, "
    "the number of OB units recorded and included, and the average number of "
    "sniff onsets and discrete ethanol contact events per trial.",
)

table1 = doc_r.add_table(rows=1, cols=6)
table1.style = "Light Grid Accent 1"
hdr = table1.rows[0].cells
headers = [
    "Session",
    "N LabView trials",
    "N units recorded",
    "N units included",
    "Avg. sniffs / trial",
    "Avg. ETH contacts / trial",
]
for cell, h in zip(hdr, headers):
    cell.text = h

for i, mt_row in enumerate(methods_table):
    sess = mt_row["session"]
    row = table1.add_row().cells
    row[0].text = str(sess)
    row[1].text = str(mt_row["n_labview_trials"])        # odd-TR LabView trials analyzed
    row[2].text = str(ui_lookup[sess]["n_units_recorded"])
    row[3].text = str(ui_lookup[sess]["n_units_included"])
    row[4].text = f"{mt_row['avg_sniffs_per_labview_trial']:.2f}"
    row[5].text = f"{mt_row['avg_eth_contacts_per_labview_trial']:.2f}"

doc_r.add_paragraph()

# ---------------------------------------------------------------------------
# 2. Main result (CON-003)
# ---------------------------------------------------------------------------
add_heading(doc_r, "Sniff Rate During Ethanol Contact vs. Control (CON-003)", level=1)

sig_phrase = "not statistically significant" if not SIG else "statistically significant"
add_para(
    doc_r,
    f"Per-animal mean instantaneous sniff rate during the first ethanol contact "
    f"period was compared to per-animal mean sniff rate during the pre-valve "
    f"control window (TS = 0-10 s) using a two-sided exact Wilcoxon signed-rank "
    f"test across n = {N_ANIMALS} animals. The test statistic was W = {W_STAT:.1f} "
    f"(exact two-sided p = {P_EXACT:.4f}; matched-pairs rank-biserial r = "
    f"{RBR:.4f}). This difference was {sig_phrase} (two-sided Wilcoxon "
    f"signed-rank test, exact) at alpha = 0.05.",
)

# ---------------------------------------------------------------------------
# 3. Per-animal means table
# ---------------------------------------------------------------------------
add_heading(doc_r, "Per-Animal Mean Sniff Rates", level=1)
add_para(
    doc_r,
    "Table 2 lists, for each animal/session, the mean instantaneous sniff rate "
    "(Hz) during the first ethanol contact period and during the pre-valve "
    "control period (TS = 0-10 s), as entered into the Wilcoxon signed-rank "
    "test above.",
)

table2 = doc_r.add_table(rows=1, cols=3)
table2.style = "Light Grid Accent 1"
hdr2 = table2.rows[0].cells
headers2 = [
    "Session",
    "Contact mean sniff rate (Hz)",
    "Control mean sniff rate (Hz)",
]
for cell, h in zip(hdr2, headers2):
    cell.text = h

for i, sess in enumerate(SESSIONS):
    row = table2.add_row().cells
    row[0].text = str(sess)
    row[1].text = f"{float(CONTACT_MEANS[i]):.4f}"
    row[2].text = f"{float(CONTROL_MEANS[i]):.4f}"

doc_r.add_paragraph()

# ---------------------------------------------------------------------------
# 4. Figure captions
# ---------------------------------------------------------------------------
add_heading(doc_r, "Figures", level=1)

# Figure 1 -- Sniff rate
add_heading(doc_r, "Figure 1: Sniff rate within each trial", level=2)
add_para(
    doc_r,
    f"Panel 1: heatmap of instantaneous sniff rate (Hz; colorbar labeled "
    f"\"sniff rate (Hz)\", jet colormap) for {N_TRIALS_FIG} LabView trials "
    f"pooled across {N_SESS_FIG} sessions, one row per trial, time from trial "
    f"start (0-40 s) on the x-axis. Panel 2: sniff rate averaged across all "
    f"trials, with a shaded band showing the 95% CI computed as mean +/- "
    f"1.96 * SEM across all {N_TRIALS_FIG} trials.",
)

# Figure 2 -- Firing rate
add_heading(doc_r, "Figure 2: Firing rate within each trial", level=2)
ui_fig2_parts = []
for sess in SESSIONS:
    n_inc = ui_lookup[sess]["n_units_included"]
    ui_fig2_parts.append(f"{sess}: {n_inc}")
add_para(
    doc_r,
    f"Panel 1: heatmap of average OB population firing rate (Hz; colorbar "
    f"labeled \"firing rate (Hz)\", jet colormap), averaged across included "
    f"units, for each of {N_TRIALS_FIG} LabView trials pooled across "
    f"{N_SESS_FIG} sessions. Units were included only if overall firing rate "
    f">= 0.1 Hz and total spike count >= 5000. Per-session included unit "
    f"counts: {'; '.join(ui_fig2_parts)}. Panel 2: population firing rate "
    f"averaged across all trials, with a shaded band showing the 95% CI "
    f"computed as mean +/- 1.96 * SEM across all {N_TRIALS_FIG} trials.",
)

# Figure 3 -- ETH
add_heading(doc_r, "Figure 3: Ethanol sensor value within each trial", level=2)
add_para(
    doc_r,
    f"Panel 1: heatmap of the raw (pre-mean-subtraction) ethanol sensor value "
    f"(colorbar labeled \"ETH sensor value (a.u., normalized 0-1)\", jet "
    f"colormap), per CON-002, for each of {N_TRIALS_FIG} LabView trials "
    f"pooled across {N_SESS_FIG} sessions. Panel 2: raw ETH sensor value "
    f"averaged across all trials, with a shaded band showing the 95% CI "
    f"computed as mean +/- 1.96 * SEM across all {N_TRIALS_FIG} trials.",
)

doc_r.save(REPORT_DIR + r"\NP-DEMO-4_results.docx")
print(f"Wrote: {REPORT_DIR}\\NP-DEMO-4_results.docx")
