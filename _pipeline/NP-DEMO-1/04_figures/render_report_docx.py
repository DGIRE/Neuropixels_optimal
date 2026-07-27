"""
render_report_docx.py — Render report_injected.md to NP-DEMO-1_ethanol_sniff_phase_locking.docx

Reads report_injected.md (all placeholders already injected) and produces a
formatted Word document using python-docx 1.2.0.

Markdown handled:
  # H1  -> Heading 1
  ## H2 -> Heading 2
  **bold** -> bold run
  `code` -> monospace run (Courier New)
  - bullet -> List Bullet style
  blank line -> paragraph break

Output: output/NP-DEMO-1_ethanol_sniff_phase_locking.docx
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

FIGURES_DIR = Path(r"C:\Projects\Repos\Neuropixels\_pipeline\NP-DEMO-1\04_figures")
OUTPUT_DIR  = FIGURES_DIR / "output"


# ---------------------------------------------------------------------------
# Inline-markup parser: bold (**...**) and code (`...`)
# ---------------------------------------------------------------------------
_INLINE_RE = re.compile(r"\*\*(.+?)\*\*|`(.+?)`")


def _add_styled_run(para, text: str, bold: bool = False, code: bool = False) -> None:
    run = para.add_run(text)
    run.bold = bold
    if code:
        run.font.name = "Courier New"
        run.font.size = Pt(9)


def _add_inline(para, text: str) -> None:
    """Add a paragraph's worth of inline markup."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            _add_styled_run(para, text[pos:m.start()])
        if m.group(1) is not None:
            _add_styled_run(para, m.group(1), bold=True)
        else:
            _add_styled_run(para, m.group(2), code=True)
        pos = m.end()
    if pos < len(text):
        _add_styled_run(para, text[pos:])


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------
def render(src: Path, dst: Path) -> None:
    lines = src.read_text(encoding="utf-8").splitlines()
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # H1
        if stripped.startswith("# ") and not stripped.startswith("## "):
            h = doc.add_heading(stripped[2:].strip(), level=1)
            i += 1

        # H2
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            i += 1

        # Bullet item
        elif stripped.startswith("- "):
            para = doc.add_paragraph(style="List Bullet")
            _add_inline(para, stripped[2:])
            i += 1

        # Blank line — skip (paragraph breaks come from add_paragraph calls)
        elif stripped == "":
            i += 1

        # Normal paragraph (may contain inline markup)
        else:
            para = doc.add_paragraph()
            _add_inline(para, stripped)
            i += 1

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(dst)
    print(f"  Saved {dst}")


def main() -> None:
    src = FIGURES_DIR / "report_injected.md"
    dst = OUTPUT_DIR / "NP-DEMO-1_ethanol_sniff_phase_locking.docx"
    render(src, dst)


if __name__ == "__main__":
    main()
